#!/usr/bin/env python3
"""Create a privacy-conscious source inventory for the PERL prototype.

The script reads the local Focused Future artifacts without modifying them. It
prints document structure, clinically relevant excerpts, and spreadsheet schema
summaries while suppressing columns likely to contain direct identifiers.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

import pandas as pd
from docx import Document
from pypdf import PdfReader


KEYWORDS = re.compile(
    r"(?i)perl|qpass|summary|counsel|clinician|indicator|diagnos|risk|"
    r"depress|anxi|trauma|substance|treatment|score|threshold|response|"
    r"valid|report|review|azure|hipaa|phi|security|pilot|university"
)
IDENTIFIER_COLUMNS = re.compile(
    r"(?i)(^|\b)(name|email|e-mail|phone|address|dob|birth|participant|"
    r"patient|client|mrn|medical record|ip address|user id)(\b|$)"
)
IDENTIFIER_HEADER_TOKENS = (
    "name",
    "email",
    "phone",
    "address",
    "dateofbirth",
    "dob",
    "birthdate",
    "participantid",
    "patientid",
    "clientid",
    "assessmentid",
    "medicalrecord",
    "mrn",
    "userid",
    "zipcode",
)
EMAIL_VALUE = re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")
PHONE_VALUE = re.compile(r"^\+?[\d\s().-]{9,}$")
TARGET_DOCX = {
    "PERL Meeting Notes .docx",
    "# QPASS Clinical Summary.docx",
    "eQPASS Features and Benefits.docx",
    "eQPASS Spec Sheet 2022 AW.docx",
    "eQPASS Validity Summary.docx",
    "Your Premium Mental Health Analysis.docx",
    "Focused Future LLC boardand team info.docx",
}
TARGET_PPTX = {"PERL_Overview.pptx", "FocusedFuture_BoardMeeting_Deck_April2026.pptx"}
TARGET_PDF = {
    "DashboardDepressionScanExample (003).pdf",
    "ScoreReportPackage-77.pdf",
    "eQPASS Spec Sheet 2022.pdf",
    "QPASS Manual Original 2007-1.PDF",
    "FocusedFuture_PPT (1).pdf",
    "Your Premium Mental Health Analysis.pdf",
}
TARGET_XLSX = {
    "WORKING COPY PERL - People Enjoy Real Living (Responses).xlsx",
    "B2C Data Export.xlsx",
}


def digest(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()[:16]


def clean(text: object, limit: int = 600) -> str:
    value = re.sub(r"\s+", " ", str(text or "")).strip()
    return value[:limit]


def relevant_lines(lines: Iterable[str], limit: int = 140) -> list[str]:
    selected: list[str] = []
    seen: set[str] = set()
    for line in lines:
        line = clean(line)
        if line and KEYWORDS.search(line) and line not in seen:
            selected.append(line)
            seen.add(line)
            if len(selected) >= limit:
                break
    return selected


def extract_docx(path: Path) -> dict:
    document = Document(path)
    paragraphs = [clean(p.text, 1200) for p in document.paragraphs if clean(p.text)]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text, 500) for cell in row.cells])
        tables.append(rows)
    return {
        "kind": "docx",
        "file": path.name,
        "sha256_16": digest(path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "relevant_paragraphs": relevant_lines(paragraphs),
        "relevant_table_rows": [
            row
            for table in tables
            for row in table
            if KEYWORDS.search(" | ".join(row))
        ][:160],
    }


def slide_number(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def extract_pptx(path: Path) -> dict:
    namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    slides = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            (
                name
                for name in archive.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ),
            key=slide_number,
        )
        for name in names:
            root = ET.fromstring(archive.read(name))
            text = [clean(node.text) for node in root.findall(".//a:t", namespace)]
            slides.append({"slide": slide_number(name), "text": [x for x in text if x]})
    return {
        "kind": "pptx",
        "file": path.name,
        "sha256_16": digest(path),
        "slide_count": len(slides),
        "slides": slides,
    }


def extract_pdf(path: Path) -> dict:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - source corruption fallback
            text = f"[text extraction failed: {type(exc).__name__}]"
        lines = relevant_lines(text.splitlines(), limit=28)
        if lines:
            pages.append({"page": index, "relevant_lines": lines})
        if len(pages) >= 18:
            break
    return {
        "kind": "pdf",
        "file": path.name,
        "sha256_16": digest(path),
        "page_count": len(reader.pages),
        "pages_with_relevant_text": pages,
    }


def summarize_series(series: pd.Series, column: str) -> dict:
    normalized_column = re.sub(r"[^a-z0-9]", "", column.lower())
    if IDENTIFIER_COLUMNS.search(column) or any(
        token in normalized_column for token in IDENTIFIER_HEADER_TOKENS
    ):
        return {"column": column, "suppressed": "possible direct identifier"}
    values = series.dropna().astype(str).map(lambda value: clean(value, 240))
    nonempty = [value for value in values if value]
    if any(EMAIL_VALUE.match(value) or PHONE_VALUE.match(value) for value in nonempty[:500]):
        return {"column": column, "suppressed": "identifier-like values detected"}
    if normalized_column.startswith(("column", "unnamed")) and nonempty:
        probable_names = sum(
            bool(re.fullmatch(r"[A-Z][A-Za-z'.-]+(?:\s+[A-Z][A-Za-z'.-]+){1,3}", value))
            for value in nonempty[:500]
        )
        if probable_names / min(len(nonempty), 500) >= 0.4:
            return {"column": column, "suppressed": "name-like values detected"}
    counts = Counter(value for value in values if value)
    return {
        "column": column,
        "non_null": int(series.notna().sum()),
        "unique": int(series.nunique(dropna=True)),
        "top_values": counts.most_common(5),
    }


def extract_xlsx(path: Path) -> dict:
    workbook = pd.ExcelFile(path)
    sheets = []
    for sheet_name in workbook.sheet_names:
        frame = pd.read_excel(workbook, sheet_name=sheet_name)
        columns = [clean(column, 240) or f"Column {index + 1}" for index, column in enumerate(frame.columns)]
        frame.columns = columns
        sheets.append(
            {
                "sheet": sheet_name,
                "rows": int(frame.shape[0]),
                "columns": int(frame.shape[1]),
                "column_summaries": [
                    summarize_series(frame[column], column) for column in columns[:120]
                ],
            }
        )
    return {
        "kind": "xlsx",
        "file": path.name,
        "sha256_16": digest(path),
        "sheets": sheets,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path)
    parser.add_argument("--only", action="append", default=[])
    args = parser.parse_args()

    extractors = {
        ".docx": (TARGET_DOCX, extract_docx),
        ".pptx": (TARGET_PPTX, extract_pptx),
        ".pdf": (TARGET_PDF, extract_pdf),
        ".xlsx": (TARGET_XLSX, extract_xlsx),
    }

    results = []
    for path in sorted(args.root.rglob("*")):
        if not path.is_file():
            continue
        if args.only and path.name not in set(args.only):
            continue
        suffix = path.suffix.lower()
        if suffix not in extractors:
            continue
        targets, extractor = extractors[suffix]
        if path.name not in targets:
            continue
        try:
            results.append(extractor(path))
        except Exception as exc:
            results.append(
                {
                    "kind": suffix.lstrip("."),
                    "file": path.name,
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )

    print(json.dumps(results, indent=2, ensure_ascii=True))


if __name__ == "__main__":
    main()
